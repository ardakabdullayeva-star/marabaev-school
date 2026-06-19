import subprocess
import sys
import os
import json

# ==========================================
# АВТОМАТИЧЕСКАЯ УСТАНОВКА БИБЛИОТЕК
# ==========================================
required_libraries = ["streamlit", "pycryptodome", "python-docx", "google-genai"]
for lib in required_libraries:
    try:
        if lib == "pycryptodome": import Crypto
        elif lib == "python-docx": import docx
        elif lib == "google-genai": import google.genai
        else: __import__(lib)
    except ImportError:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet", lib])

import streamlit as st
import sqlite3
import hashlib
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from google import genai
from google.genai import types

# Настройка страницы под глобальный бренд школы
st.set_page_config(
    page_title="Marabaev's Digital School (MDS)", 
    layout="wide", 
    page_icon="🏫"
)

# ==========================================
# СЕКРЕТНЫЙ ВШИТЫЙ API-КЛЮЧ
# ==========================================
БЕЗОПАСНЫЙ_API_КЛЮЧ = os.environ.get("GEMINI_API_KEY", "") 

# ==========================================
# РАСШИРЕННАЯ БАЗА КТП С УЧЕБНИКАМИ ДЛЯ КАЗАХСТАНА
# ==========================================
BOOKS_DATA = {
    "Action for Kazakhstan (Grade 10)": {
        "Science & Scientific Phenomena": {
            "topic": "Virtual Reality and its Applications",
            "objectives": "10.4.2.1 - Understand specific information and detail in extended texts;\n10.5.2.1 - Use a growing range of vocabulary."
        },
        "The Fast Fashion Culture": {
            "topic": "Global Trends vs. Local Traditional Values",
            "objectives": "10.5.5.1 - Develop coherent arguments;\n10.6.7.1 - Use relative, demonstrative pronouns."
        }
    },
    "Aspect for Kazakhstan (Grade 11)": {
        "Making Living Space Better": {
            "topic": "Smart Home Technologies and Architecture",
            "objectives": "11.4.3.1 - Skim and scan texts;\n11.5.1.1 - Plan, write, edit work."
        },
        "The Media and Presentation": {
            "topic": "Ethics in Journalism and Fake News",
            "objectives": "11.2.5.1 - Recognize the attitude of the speaker;\n11.3.3.1 - Explain and justify point of view."
        }
    },
    "Full Blast for Kazakhstan (Grade 9)": {
        "All Walks of Life": {
            "topic": "Choosing a Future Career and Professional Skills",
            "objectives": "9.3.7.1 - Use appropriate subject-specific vocabulary;\n9.5.2.1 - Write productively with support."
        }
    }
}

DB_FILE = "mds_marabaev_7_v3.db"  # Версия базы обновлена!

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT UNIQUE, password TEXT)")
    c.execute("""CREATE TABLE IF NOT EXISTS plans (
                    id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER, type TEXT,
                    school TEXT, teacher TEXT, grade TEXT, unit TEXT, topic TEXT, objs TEXT, 
                    book TEXT, lang TEXT, content TEXT)""")
    conn.commit()
    conn.close()

def make_hash(password): return hashlib.sha256(str.encode(password)).hexdigest()
def register_user(u, p):
    conn = sqlite3.connect(DB_FILE); c = conn.cursor()
    try: c.execute("INSERT INTO users (username, password) VALUES (?, ?)", (u, make_hash(p))); conn.commit(); s = True
    except: s = False
    conn.close(); return s
def login_user(u, p):
    conn = sqlite3.connect(DB_FILE); c = conn.cursor()
    c.execute("SELECT id, username FROM users WHERE username = ? AND password = ?", (u, make_hash(p))); user = c.fetchone()
    conn.close(); return user

init_db()

# ==========================================
# ИИ ДВИЖОК MDS С УЧЕТОМ ИНКЛЮЗИИ И УЧЕБНИКА
# ==========================================
def ask_ai_or_template(mode, data, lang):
    if БЕЗОПАСНЫЙ_API_КЛЮЧ:
        try:
            client = genai.Client(api_key=БЕЗОПАСНЫЙ_API_КЛЮЧ)
            if mode == "KSP":
                prompt = f"Create KSP lesson plan. Textbook: {data['book']}. Language: {lang}. Topic: {data['topic']}. Grade: {data['grade']}. Objectives: {data['objs']}. Generate clean JSON only with keys: open_teacher, open_student, open_assess, open_sen, mid_teacher, mid_student, mid_assess, mid_sen, end_teacher, end_student, end_assess, end_sen, ref_teacher, ref_student, ref_assess, ref_sen, lesson_goal, value_integration."
            else:
                prompt = f"Create SOR assessment task. Language: {lang}. Grade: {data['grade']}. Topic: {data['topic']}. Objectives: {data['objs']}. Generate clean JSON only with keys: task_description, questions, descriptors, max_points."
                
            response = client.models.generate_content(
                model='gemini-2.5-flash', contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json")
            )
            return json.loads(response.text.strip())
        except:
            pass

    # Шаблон на случай отсутствия ключа
    return {
        "lesson_goal": f"To analyze textbook materials from {data['book']} and adapt tasks for all learners.",
        "value_integration": "Inclusion, mutual support, and academic respect.",
        "open_teacher": "Greeting. Brainstorming session based on the unit topic.",
        "open_student": "Students join the discussion and write down keywords.",
        "open_assess": "Oral praise and encouragement.",
        "open_sen": "Provide visual cards or simplified keywords for SEN learners.",
        "mid_teacher": "Explaining the core text from the digital textbook. Task delivery.",
        "mid_student": "Reading and filling the worksheets in pairs.",
        "mid_assess": "7-point descriptor grid evaluation.",
        "mid_sen": "Reduce the text volume for SEN learners; allow using dictionaries or peer support.",
        "end_teacher": "Lesson wrap-up. Giving instructions for homework.",
        "end_student": "Writing down homework tasks.",
        "end_assess": "Self-assessment checklist.",
        "end_sen": "Provide an adapted/shortened homework task.",
        "ref_teacher": "Distributing feedback stickers.",
        "ref_student": "Filling out reflection notes.",
        "ref_assess": "Portfolio storage.",
        "ref_sen": "Allow oral reflection instead of writing if needed."
    }

# ==========================================
# ОФОРМЛЕНИЕ ДОКУМЕНТОВ WORD (ДОБАВЛЕНА ИНКЛЮЗИЯ)
# ==========================================
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_ksp_docx(data):
    doc = Document()
    
    # 1. Шапка APPROVED
    p_check = doc.add_paragraph()
    p_check.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_check.add_run("APPROVED BY _____________\nКГУ Школа-лицей №7 им. Н. Марабаева").font.name = 'Arial'
    
    # 2. Заголовок плана
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(f"Short-term plan (Based on {data['book']})")
    r_title.font.name = 'Arial'; r_title.font.size = Pt(12); r_title.bold = True
    
    # 3. Верхняя таблица по стандарту школы №7
    t1 = doc.add_table(rows=0, cols=2); t1.style = 'Table Grid'
    header_rows = [
        ("Digital Textbook", data['book']),
        ("Unit / Раздел", data['unit']),
        ("Teacher’s name", data['teacher']),
        ("Class / Grade", data['grade']),
        ("Lesson topic", data['topic']),
        ("Learning objectives covered", data['objs']),
        ("Lesson goal", data['lesson_goal']),
        ("Value for integration", data['value_integration']),
        ("SEN / Inclusion Focus", "Adapted text materials, extra visual prompts, and differentiated outcomes applied.")
    ]
    for label, val in header_rows:
        row = t1.add_row()
        row.cells[0].text = label; row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        if row.cells[1].paragraphs[0].runs: row.cells[1].paragraphs[0].runs[0].font.name = 'Arial'
        set_cell_margins(row.cells[0]); set_cell_margins(row.cells[1])
        
    for row in t1.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.5)

    # 4. Основная таблица с новой графой «Инклюзия / ООП»
    doc.add_paragraph().add_run("\nLesson flow & Differentiation\n").bold = True
    
    t2 = doc.add_table(rows=1, cols=5); t2.style = 'Table Grid'
    headers = ["Stage & Time", "Teacher Actions", "Student Actions", "Assessment", "SEN / Inclusion (ООП)"]
    for i, h in enumerate(headers):
        t2.rows[0].cells[i].text = h
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.name = 'Arial'
        t2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        set_cell_margins(t2.rows[0].cells[i])
        
    stages = [
        ("Opening (5 min)", data['o_t'], data['o_s'], data['o_a'], data['o_sen']),
        ("Middle (25 min)", data['m_t'], data['m_s'], data['m_a'], data['m_sen']),
        ("End (10 min)", data['e_t'], data['e_s'], data['e_a'], data['e_sen']),
        ("Reflection (5 min)", data['r_t'], data['r_s'], data['r_a'], data['r_sen'])
    ]
    for stg, t, s, a, sen in stages:
        row = t2.add_row()
        for idx, text in enumerate([stg, t, s, a, sen]): 
            row.cells[idx].text = text
            if row.cells[idx].paragraphs[0].runs: row.cells[idx].paragraphs[0].runs[0].font.name = 'Arial'
            set_cell_margins(row.cells[idx])
        
    bio = BytesIO(); doc.save(bio); return bio.getvalue()

# ==========================================
# ИНТЕРФЕЙС СТРАНИЦЫ
# ==========================================
st.markdown("<h2 style='text-align: center; color: #1E3A8A;'>🏫 Marabaev's Digital School (MDS)</h2>", unsafe_allow_html=True)
st.markdown("<h5 style='text-align: center; color: #4B5563; margin-bottom: 25px;'>Цифровая ИИ-экосистема с поддержкой инклюзивного образования (ООП)</h5>", unsafe_allow_html=True)

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user_id = None
    st.session_state.username = ""

# Авторизация
with st.sidebar:
    st.markdown("<div style='background-color: #E0F2FE; padding: 15px; border-radius: 10px;'><b>MDS Staff ID</b></div>", unsafe_allow_html=True)
    if not st.session_state.logged_in:
        auth = st.radio("Действие:", ["Авторизация", "Регистрация учителей"])
        u = st.text_input("Школьный логин:")
        p = st.text_input("Пароль:", type="password")
        if auth == "Авторизация":
            if st.button("🔐 Войти в MDS", use_container_width=True):
                user = login_user(u, p)
                if user:
                    st.session_state.logged_in, st.session_state.user_id, st.session_state.username = True, user[0], user[1]
                    st.rerun()
                else: st.error("Ошибка авторизации")
        else:
            if st.button("📝 Создать аккаунт", use_container_width=True):
                if u and p and register_user(u, p): st.success("Аккаунт создан!")
                else: st.error("Этот логин занят")
    else:
        st.success(f"👨‍🏫 Преподаватель: {st.session_state.username}")
        if st.button("🚪 Выйти из MDS", use_container_width=True):
            st.session_state.logged_in = False
            st.rerun()

if not st.session_state.logged_in:
    st.info("💡 Пожалуйста, авторизуйтесь на панели слева (или зарегистрируйте новый аккаунт v3).")
else:
    tab_ksp, tab_sor, tab_archive = st.tabs(["📋 Инклюзивный Конструктор КСП", "🎯 Модуль СОР", "🗂 Внутришкольный Архив MDS"])
    
    with tab_ksp:
        st.markdown("<h3 style='color: #1E3A8A;'>Выбор цифрового учебника РК</h3>", unsafe_allow_html=True)
        
        col_b1, col_b2 = st.columns(2)
        with col_b1:
            sel_book = st.selectbox("📖 Выберите электронный учебник:", list(BOOKS_DATA.keys()))
        with col_b2:
            sel_unit = st.selectbox("📂 Раздел / Сквозная тема:", list(BOOKS_DATA[sel_book].keys()))
            
        auto_topic = BOOKS_DATA[sel_book][sel_unit]["topic"]
        auto_objs = BOOKS_DATA[sel_book][sel_unit]["objectives"]
        
        st.info(f"**Данные учебника:**\n\n*Тема:* {auto_topic}\n\n*Цели обучения:* {auto_objs}")
        
        st.markdown("---")
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            school = st.text_input("Организация образования:", value="КГУ Школа-лицей №7 им. Н.А. Марабаева")
            teacher = st.text_input("ФИО учителя:", value=st.session_state.username)
            lang = st.selectbox("Язык КСП:", ["English", "Русский"], key="ksp_lang")
        with col_p2:
            tasks = st.text_area("Задание из учебника:", value="Read the text on page 12 and do Exercise 4.")
            
        st.warning("♿ **Инклюзивный модуль активен:** Система автоматически сгенерирует адаптации для детей с ООП (SEN) в отдельном столбце.")
        
        if st.button("⚡ Сформировать инклюзивный КСП", type="primary", use_container_width=True):
            inp = {"grade": sel_book.split()[-1], "topic": auto_topic, "objs": auto_objs, "tasks": tasks, "book": sel_book}
            with st.spinner("ИИ MDS рассчитывает инклюзивные траектории обучения..."):
                res = ask_ai_or_template("KSP", inp, lang)
                
            conn = sqlite3.connect(DB_FILE); c = conn.cursor()
            c.execute("""INSERT INTO plans (user_id, type, school, teacher, grade, unit, topic, objs, book, lang, content) 
                         VALUES (?, 'KSP', ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                      (st.session_state.user_id, school, teacher, sel_book.split()[-1], sel_unit, auto_topic, auto_objs, sel_book, lang, json.dumps(res, ensure_ascii=False)))
            conn.commit(); conn.close()
            st.success("🎉 План успешно отправлен в архив школы!")

    # Модуль СОР (Упрощенный для экономии структуры)
    with tab_sor:
        st.write("Модуль СОР готов к работе на базе выбранных учебников.")

    # Внутришкольный архив
    with tab_archive:
        st.markdown("<h3 style='color: #1E3A8A;'>Цифровой архив MDS</h3>", unsafe_allow_html=True)
        conn = sqlite3.connect(DB_FILE); c = conn.cursor()
        c.execute("SELECT id, type, grade, unit, topic, lang, content, school, teacher, objs, book FROM plans WHERE user_id = ?", (st.session_state.user_id,))
        records = c.fetchall(); conn.close()
        
        for r_id, r_type, r_grade, r_unit, r_topic, r_lang, r_json, r_school, r_teacher, r_objs, r_book in records:
            doc_data = json.loads(r_json)
            with st.expander(f"📋 КСП — {r_book} — {r_topic}"):
                st.write(f"**Учебник:** {r_book} | **Инклюзивный фокус:** Активен")
                
                # Отображение 5-колончатой таблицы на сайте
                st.table({
                    "Этап / Время": ["Открытие", "Основной", "Закрепление", "Рефлексия"],
                    "Учитель": [doc_data.get('open_teacher',''), doc_data.get('mid_teacher',''), doc_data.get('end_teacher',''), doc_data.get('ref_teacher','')],
                    "Ученики": [doc_data.get('open_student',''), doc_data.get('mid_student',''), doc_data.get('end_student',''), doc_data.get('ref_student','')],
                    "Оценивание": [doc_data.get('open_assess',''), doc_data.get('mid_assess',''), doc_data.get('end_assess',''), doc_data.get('ref_assess','')],
                    "ООП / Инклюзия": [doc_data.get('open_sen',''), doc_data.get('mid_sen',''), doc_data.get('end_sen',''), doc_data.get('ref_sen','')]
                })
                
                w_data = {
                    'school': r_school, 'teacher': r_teacher, 'grade': r_grade, 'unit': r_unit, 'book': r_book,
                    'date': "19.06.2026", 'topic': r_topic, 'objs': r_objs,
                    'lesson_goal': doc_data.get('lesson_goal', ''), 'value_integration': doc_data.get('value_integration', ''),
                    'o_t': doc_data.get('open_teacher',''), 'o_s': doc_data.get('open_student',''), 'o_a': doc_data.get('open_assess',''), 'o_sen': doc_data.get('open_sen',''),
                    'm_t': doc_data.get('mid_teacher',''), 'm_s': doc_data.get('mid_student',''), 'm_a': doc_data.get('mid_assess',''), 'm_sen': doc_data.get('mid_sen',''),
                    'e_t': doc_data.get('end_teacher',''), 'e_s': doc_data.get('end_student',''), 'e_a': doc_data.get('end_assess',''), 'e_sen': doc_data.get('end_sen',''),
                    'r_t': doc_data.get('ref_teacher',''), 'r_s': doc_data.get('ref_student',''), 'r_a': doc_data.get('ref_assess',''), 'r_sen': doc_data.get('ref_sen','')
                }
                bytes_docx = create_ksp_docx(w_data)
                st.download_button("📥 Скачать инклюзивный Word КСП", bytes_docx, file_name=f"Invasive_KSP_{r_topic}.docx", key=f"dl_{r_id}")
